"""
自动化生成docx脚本 —— 按《雪山降临》格式
支持：
  - 从结构化数据生成
  - 从文本标记生成
  - 批量读取txt文件夹，自动识别对话/旁白，生成带封面+目录的完整书籍

字体统一策略：
  所有 run 和 style 均通过 _set_run_font() / _FONT_XML 同时设置
  w:ascii + w:hAnsi + w:eastAsia + w:cs，杜绝 Word fallback 到 Calibri。
"""

import os
import re
import docx
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

# ---- 全局默认字体名 ----
_FONT = '等线'

# 预构建的 rFonts XML 片段（避免到处字符串拼接）
_FONTS_XML = f'<w:rFonts {nsdecls("w")} w:ascii="{_FONT}" w:hAnsi="{_FONT}" w:eastAsia="{_FONT}" w:cs="{_FONT}"/>'


# =====================================================================
# 字体统一工具
# =====================================================================

def _set_run_font(run, font_name=_FONT):
    """在 run 的 rPr 中写入或覆盖 <w:rFonts>，四项字体全部设为 font_name。"""
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        run._r.insert(0, rPr)
    # 移除旧的 rFonts
    old = rPr.find(qn('w:rFonts'))
    if old is not None:
        rPr.remove(old)
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} '
        f'w:ascii="{font_name}" w:hAnsi="{font_name}" '
        f'w:eastAsia="{font_name}" w:cs="{font_name}"/>')
    rPr.insert(0, rFonts)


def _add_run_with_font(para, text, font_name=_FONT, **font_kw):
    """在段落中添加一个已设置好四项字体的 run。font_kw 可传 size, color, bold 等。"""
    run = para.add_run(text)
    _set_run_font(run, font_name)
    if 'size' in font_kw:
        run.font.size = font_kw['size']
    if 'color' in font_kw:
        run.font.color.rgb = font_kw['color']
    if font_kw.get('bold') is True:
        run.font.bold = True
    if font_kw.get('italic') is True:
        run.font.italic = True
    return run


def _ensure_style_rfonts(style, font_name=_FONT):
    """确保样式有完整的四项字体 rFonts 元素，覆盖旧的不完整设置。"""
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        style.element.insert(0, rPr)
    old = rPr.find(qn('w:rFonts'))
    if old is not None:
        rPr.remove(old)
    rPr.insert(0, parse_xml(
        f'<w:rFonts {nsdecls("w")} '
        f'w:ascii="{font_name}" w:hAnsi="{font_name}" '
        f'w:eastAsia="{font_name}" w:cs="{font_name}"/>'))


# =====================================================================
# 文档创建 & 样式设置
# =====================================================================

def create_document():
    """创建并返回一个匹配《雪山降临》格式的空白docx文档"""
    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(3.175)
    section.right_margin  = Cm(3.175)
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # ---- docDefaults ----
    _setup_doc_defaults(doc)

    # ---- Normal ----
    _setup_normal(doc)

    # ---- 各样式 ----
    _setup_heading2(doc)
    _setup_intense_quote(doc)
    _setup_quote(doc)
    _setup_title(doc)
    _setup_toc_heading(doc)

    # ---- 主题字体 ----
    _set_theme_fonts(doc)

    return doc


def _setup_doc_defaults(doc):
    """docDefaults：用显式 等线 替代 theme 引用，四项字体全部设齐。"""
    styles_xml = doc.styles.element
    defaults = styles_xml.find(qn('w:docDefaults'))
    if defaults is None:
        defaults = parse_xml(f'<w:docDefaults {nsdecls("w")}></w:docDefaults>')
        styles_xml.insert(0, defaults)

    # rPrDefault
    rPrDefault = defaults.find(qn('w:rPrDefault'))
    if rPrDefault is None:
        rPrDefault = parse_xml(f'<w:rPrDefault {nsdecls("w")}><w:rPr></w:rPr></w:rPrDefault>')
        defaults.insert(0, rPrDefault)
    rPr = rPrDefault.find(qn('w:rPr'))
    for child in list(rPr):
        rPr.remove(child)
    rPr.append(parse_xml(_FONTS_XML))
    rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="22"/>'))
    rPr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="24"/>'))
    rPr.append(parse_xml(f'<w:kern {nsdecls("w")} w:val="2"/>'))
    rPr.append(parse_xml(
        f'<w:lang {nsdecls("w")} w:val="en-US" w:eastAsia="zh-CN" w:bidi="ar-SA"/>'))

    # pPrDefault
    pPrDefault = defaults.find(qn('w:pPrDefault'))
    if pPrDefault is None:
        pPrDefault = parse_xml(f'<w:pPrDefault {nsdecls("w")}><w:pPr></w:pPr></w:pPrDefault>')
        defaults.insert(1, pPrDefault)
    pPr = pPrDefault.find(qn('w:pPr'))
    for child in list(pPr):
        pPr.remove(child)
    pPr.append(parse_xml(
        f'<w:spacing {nsdecls("w")} w:after="160" w:line="278" w:lineRule="auto"/>'))


def _setup_normal(doc):
    """Normal 样式：显式等线四项字体、11pt、黑色、无窗口控制。"""
    normal = doc.styles['Normal']
    for child in list(normal.element):
        tag = _local_tag(child)
        if tag in ('pPr', 'rPr'):
            normal.element.remove(child)

    pPr = parse_xml(f'<w:pPr {nsdecls("w")}><w:widowControl w:val="0"/></w:pPr>')
    normal.element.insert(0, pPr)

    rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
    normal.element.insert(0, rPr)
    rPr.insert(0, parse_xml(_FONTS_XML))
    rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="22"/>'))


def _setup_heading2(doc):
    h2 = doc.styles['Heading 2']
    for child in list(h2.element):
        tag = _local_tag(child)
        if tag in ('pPr', 'rPr'):
            h2.element.remove(child)

    pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
    h2.element.insert(0, pPr)
    pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))
    pPr.append(parse_xml(f'<w:keepLines {nsdecls("w")}/>'))
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="160" w:after="80"/>'))
    pPr.append(parse_xml(f'<w:outlineLvl {nsdecls("w")} w:val="1"/>'))

    rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
    h2.element.insert(0, rPr)
    rPr.insert(0, parse_xml(_FONTS_XML))
    rPr.append(parse_xml(
        f'<w:color {nsdecls("w")} w:val="0000FA"/>'))
    rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="40"/>'))
    rPr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="40"/>'))
    rPr.append(parse_xml(f'<w:b {nsdecls("w")} w:val="false"/>'))
    rPr.append(parse_xml(f'<w:bCs {nsdecls("w")} w:val="false"/>'))


def _setup_intense_quote(doc):
    iq = doc.styles['Intense Quote']
    for child in list(iq.element):
        tag = _local_tag(child)
        if tag in ('pPr', 'rPr'):
            iq.element.remove(child)

    pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
    iq.element.insert(0, pPr)
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="10" '
        f'w:color="0000FA"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="10" '
        f'w:color="0000FA"/>'
        f'</w:pBdr>'))
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="360" w:after="360"/>'))
    pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="864" w:right="864"/>'))
    pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))

    rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
    iq.element.insert(0, rPr)
    rPr.insert(0, parse_xml(_FONTS_XML))
    rPr.append(parse_xml(f'<w:i {nsdecls("w")}/>'))
    rPr.append(parse_xml(f'<w:iCs {nsdecls("w")}/>'))
    rPr.append(parse_xml(
        f'<w:color {nsdecls("w")} w:val="0000FA"/>'))


def _setup_quote(doc):
    q = doc.styles['Quote']
    for child in list(q.element):
        tag = _local_tag(child)
        if tag in ('pPr', 'rPr'):
            q.element.remove(child)

    pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
    q.element.insert(0, pPr)
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="160"/>'))
    pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))

    rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
    q.element.insert(0, rPr)
    rPr.insert(0, parse_xml(_FONTS_XML))
    rPr.append(parse_xml(f'<w:i {nsdecls("w")}/>'))
    rPr.append(parse_xml(f'<w:iCs {nsdecls("w")}/>'))
    rPr.append(parse_xml(
        f'<w:color {nsdecls("w")} w:val="404040" w:themeColor="text1" w:themeTint="BF"/>'))


def _setup_title(doc):
    title = doc.styles['Title']
    for child in list(title.element):
        tag = _local_tag(child)
        if tag in ('pPr', 'rPr'):
            title.element.remove(child)

    pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
    title.element.insert(0, pPr)
    pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:line="480" w:lineRule="auto"/>'))

    rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
    title.element.insert(0, rPr)
    rPr.insert(0, parse_xml(_FONTS_XML))
    rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="56"/>'))
    rPr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="56"/>'))
    rPr.append(parse_xml(
        f'<w:color {nsdecls("w")} w:val="0000FA"/>'))


def _setup_toc_heading(doc):
    toc = doc.styles['TOC Heading']
    for child in list(toc.element):
        tag = _local_tag(child)
        if tag in ('pPr', 'rPr'):
            toc.element.remove(child)

    pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
    toc.element.insert(0, pPr)

    rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
    toc.element.insert(0, rPr)
    rPr.insert(0, parse_xml(_FONTS_XML))
    rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="32"/>'))
    rPr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="32"/>'))
    rPr.append(parse_xml(
        f'<w:color {nsdecls("w")} w:val="0000FA"/>'))


def _set_theme_fonts(doc):
    """将主题 major/minor 字体也统一设为等线族，兜底防止 theme fallback。"""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    try:
        theme_part = doc.part.part_related_by(RT.THEME)
        for el in theme_part.element.findall('.//' + qn('a:majorFont') + '/' + qn('a:latin')):
            el.set('typeface', '等线')
        for el in theme_part.element.findall('.//' + qn('a:majorFont') + '/' + qn('a:ea')):
            el.set('typeface', '等线')
        for el in theme_part.element.findall('.//' + qn('a:minorFont') + '/' + qn('a:latin')):
            el.set('typeface', '等线')
        for el in theme_part.element.findall('.//' + qn('a:minorFont') + '/' + qn('a:ea')):
            el.set('typeface', '等线')
    except Exception:
        pass


def _local_tag(el):
    tag = el.tag
    return tag.split('}')[-1] if '}' in tag else tag


# =====================================================================
# TXT 解析
# =====================================================================

def parse_txt_file(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    content = []
    for line in lines:
        line = line.rstrip('\n').rstrip('\r')
        if not line.strip():
            continue
        if line.startswith('#'):
            continue
        match = _match_dialogue(line)
        if match:
            content.append({"type": "dialogue", "speaker": match[0], "text": match[1]})
        elif line.strip():
            content.append({"type": "narrative", "text": line.strip()})
    return content


def _match_dialogue(line):
    if ':' not in line:
        return None
    parts = line.split(':', 1)
    speaker = parts[0].strip()
    text = parts[1].strip() if len(parts) > 1 else ''
    if not speaker or len(speaker) > 12:
        return None
    if speaker.isdigit():
        return None
    if speaker[0].isdigit():
        return None
    for ch in speaker:
        cp = ord(ch)
        if (0x2E80 <= cp <= 0x9FFF or
            0xF900 <= cp <= 0xFAFF or
            0xFE30 <= cp <= 0xFE4F or
            0xFF00 <= cp <= 0xFFEF or
            0x20000 <= cp <= 0x2A6DF or
            0x2F800 <= cp <= 0x2FA1F):
            return (speaker, text)
    if all(c.isascii() and (c.isalpha() or c in "'.- ") for c in speaker) and speaker[0].isalpha():
        return (speaker, text)
    return None


# =====================================================================
# 封面 & 目录
# =====================================================================

def _apply_duotone(para):
    blips = para._p.findall('.//' + qn('a:blip'))
    for blip in blips:
        old_duo = blip.find(qn('a:duotone'))
        if old_duo is not None:
            blip.remove(old_duo)
        duotone = parse_xml(
            f'<a:duotone {nsdecls("a")}>'
            f'<a:srgbClr val="0000FA"/>'
            f'<a:prstClr val="white"/>'
            f'</a:duotone>')
        blip.append(duotone)


def add_cover(doc, title, subtitle='明日方舟'):
    from docx.enum.text import WD_BREAK

    cover_img_paths = [
        '_cover_assets/image1.png',
        os.path.join(os.path.dirname(__file__), '_cover_assets/image1.png'),
    ]
    img_path = None
    for p in cover_img_paths:
        if os.path.isfile(p):
            img_path = p
            break

    if img_path:
        para_img = doc.add_paragraph(style='No Spacing')
        para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = para_img._p.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            para_img._p.insert(0, pPr)
        pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="1540" w:after="240"/>'))
        run = para_img.add_run()
        run.add_picture(img_path, width=docx.shared.Cm(5))
        _apply_duotone(para_img)
    else:
        for _ in range(3):
            doc.add_paragraph(style='No Spacing')

    # 主标题
    para_title = doc.add_paragraph(style='No Spacing')
    para_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = para_title._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
        para_title._p.insert(0, pPr)
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="6" w:space="10" w:color="0000FA"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="10" w:color="0000FA"/>'
        f'</w:pBdr>'))
    pPr.append(parse_xml(f'<w:jc {nsdecls("w")} w:val="center"/>'))

    run_title = _add_run_with_font(para_title, title,
                                   size=Pt(40), color=RGBColor(0x00, 0x00, 0xFA))

    # 副标题
    doc.add_paragraph(style='No Spacing').text = ''
    para_sub = doc.add_paragraph(style='No Spacing')
    para_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run_with_font(para_sub, subtitle,
                       size=Pt(16), color=RGBColor(0x00, 0x00, 0xFA))


    doc.add_paragraph(style='No Spacing')

    # 底部图片
    img2_path = None
    for p in ['_cover_assets/image2.png',
              os.path.join(os.path.dirname(__file__), '_cover_assets/image2.png')]:
        if os.path.isfile(p):
            img2_path = p
            break
    if img2_path:
        para_img2 = doc.add_paragraph(style='No Spacing')
        para_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = para_img2.add_run()
        run2.add_picture(img2_path, width=docx.shared.Cm(3))
        _apply_duotone(para_img2)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_toc(doc, entries):
    toc_title = doc.add_paragraph('目录', style='TOC Heading')

    doc.add_paragraph(style='Normal').text = ''

    for i, entry in enumerate(entries):
        if isinstance(entry, (list, tuple)):
            name, page = entry[0], entry[1]
        else:
            name, page = entry, str(i + 1)

        para = doc.add_paragraph(style='Normal')

        # 制表位
        pPr = para._p.find(qn('w:pPr'))
        if pPr is None:
            pPr = parse_xml(f'<w:pPr {nsdecls("w")}></w:pPr>')
            para._p.insert(0, pPr)
        pPr.append(parse_xml(
            f'<w:tabs {nsdecls("w")}>'
            f'<w:tab w:val="right" w:leader="dot" w:pos="8296"/>'
            f'</w:tabs>'))
        pPr.append(parse_xml(f'<w:ind {nsdecls("w")} w:left="440"/>'))

        _add_run_with_font(para, name, size=Pt(12))
        _add_run_with_font(para, '\t')
        _add_run_with_font(para, str(page), size=Pt(12))

    doc.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.PAGE)


# =====================================================================
# 便捷添加函数（所有 run 均通过 _add_run_with_font 确保字体统一）
# =====================================================================

def add_chapter_title(doc, title_text):
    doc.add_paragraph(title_text, style='Heading 2')


def add_dialogue(doc, speaker, text):
    para = doc.add_paragraph(style='Normal')
    _add_run_with_font(para, f'{speaker}：{text}')


def add_narrative(doc, text):
    doc.add_paragraph(text, style='Intense Quote')


def add_quote(doc, text):
    doc.add_paragraph(text, style='Quote')


def add_separator(doc, text='战'):
    add_quote(doc, text)


def add_normal(doc, text):
    para = doc.add_paragraph(style='Normal')
    _add_run_with_font(para, text)


# =====================================================================
# 单章节生成
# =====================================================================

def add_chapter_from_content(doc, title, content_items):
    doc.add_page_break()
    add_chapter_title(doc, title)
    for item in content_items:
        if item['type'] == 'dialogue':
            add_dialogue(doc, item['speaker'], item['text'])
        elif item['type'] == 'narrative':
            add_narrative(doc, item['text'])


# =====================================================================
# 批量生成（核心功能）
# =====================================================================

def batch_generate(folder_path, output_path=None):
    folder_path = os.path.abspath(folder_path)
    folder_name = os.path.basename(folder_path)

    if output_path is None:
        output_path = r"E:\czh要的程序\export"

    txt_files = []
    for f in os.listdir(folder_path):
        if f.endswith('.txt') and not f.startswith('~'):
            txt_files.append(f)

    if not txt_files:
        print(f'错误: 文件夹 {folder_path} 中没有找到txt文件')
        return None

    txt_files.sort(key=_sort_key)

    print(f'找到 {len(txt_files)} 个txt文件:')
    for f in txt_files:
        print(f'  {f}')

    all_chapters = []
    for filename in txt_files:
        filepath = os.path.join(folder_path, filename)
        chapter_title = os.path.splitext(filename)[0]
        content = parse_txt_file(filepath)
        all_chapters.append({'title': chapter_title, 'content': content})

    page_numbers, toc_pages = _calc_page_numbers(all_chapters)
    content_start_page = 1 + toc_pages

    doc = create_document()

    add_cover(doc, folder_name)

    toc_entries = []
    for i, ch in enumerate(all_chapters):
        toc_entries.append((ch['title'], content_start_page + page_numbers[i]))
    add_toc(doc, toc_entries)

    for ch in all_chapters:
        print(f'  处理: {ch["title"]} ...')
        add_chapter_from_content(doc, ch['title'], ch['content'])

    doc.save(output_path)
    print(f'\n已生成: {output_path}')
    print(f'  封面: {folder_name}')
    print(f'  目录: {len(toc_entries)} 章')
    print(f'  总字数: 约 {_count_chars(doc)} 字')
    return output_path


def _calc_page_numbers(all_chapters):
    PAGE_HEIGHT_PT = 697
    page_offsets = []
    cumulative = 0
    for ch in all_chapters:
        page_offsets.append(cumulative)
        h = 32
        for item in ch['content']:
            text = item.get('text', '')
            if item['type'] == 'dialogue':
                full_text = f"{item.get('speaker', '')}：{text}"
                lines = max(1, len(full_text) / 38)
                h += lines * 20.7
            elif item['type'] == 'narrative':
                lines = max(1, len(text) / 38)
                h += 48 + lines * 12.7
        cumulative += h
    page_offsets = [int(offset / PAGE_HEIGHT_PT) for offset in page_offsets]
    toc_pages = max(1, (len(all_chapters) + 33) // 34)
    return page_offsets, toc_pages


def _sort_key(filename):
    name = os.path.splitext(filename)[0]
    m = re.match(r'(\d+)-(\d+)\s', name)
    if m:
        major = int(m.group(1))
        minor = int(m.group(2))
        has_end = 1 if '后' in name or 'END' in name.upper() else 0
        return (major, minor, has_end, name)
    m2 = re.match(r'(\d+)', name)
    if m2:
        return (int(m2.group(1)), 0, 0, name)
    return (9999, 0, 0, name)


def _count_chars(doc):
    return sum(len(p.text) for p in doc.paragraphs)


# =====================================================================
# JSON / 文本标记接口（保留）
# =====================================================================

def generate_from_data(data, output_path):
    doc = create_document()
    if isinstance(data, dict):
        data = [data]
    for chapter in data:
        title = chapter.get('title', '')
        if title:
            add_chapter_title(doc, title)
        for item in chapter.get('content', []):
            if isinstance(item, str) and os.path.isfile(item):
                for p in parse_txt_file(item):
                    _add_item(doc, p)
                continue
            _add_item(doc, item, item.get('type', ''))
    doc.save(output_path)
    print(f'已生成: {output_path}')
    return output_path


def _add_item(doc, item, t=None):
    if t is None:
        t = item.get('type', '')
    if t == 'dialogue':
        add_dialogue(doc, item.get('speaker', ''), item.get('text', ''))
    elif t == 'narrative':
        add_narrative(doc, item.get('text', ''))
    elif t == 'separator':
        add_separator(doc, item.get('text', '战'))
    elif t == 'quote':
        add_quote(doc, item.get('text', ''))
    else:
        add_normal(doc, str(item.get('text', item)))


def generate_from_text(text, output_path):
    doc = create_document()
    for line in text.strip().split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            add_chapter_title(doc, line[2:])
        elif line.startswith('> '):
            add_narrative(doc, line[2:])
        elif line == '---':
            add_separator(doc)
        elif '：' in line or ':' in line:
            sep = '：' if '：' in line else ':'
            parts = line.split(sep, 1)
            add_dialogue(doc, parts[0], parts[1] if len(parts) > 1 else '')
        else:
            add_normal(doc, line)
    doc.save(output_path)
    print(f'已生成: {output_path}')
    return output_path


# =====================================================================
# 命令行入口
# =====================================================================

if __name__ == '__main__':
    import argparse, json

    parser = argparse.ArgumentParser(description='按《雪山降临》格式生成docx文档')
    parser.add_argument('--folder', '-f', help='批量模式: 包含txt文件的文件夹路径')
    parser.add_argument('--json', '-j', help='JSON数据文件路径')
    parser.add_argument('--text', '-t', help='文本文件路径（markdown风格标记）')
    parser.add_argument('--output', '-o', default=None, help='输出文件路径')

    args = parser.parse_args()

    if args.folder:
        batch_generate(args.folder, args.output)
    elif args.json:
        with open(args.json, 'r', encoding='utf-8') as f:
            data = json.load(f)
        generate_from_data(data, args.output or 'output.docx')
    elif args.text:
        with open(args.text, 'r', encoding='utf-8') as f:
            text = f.read()
        generate_from_text(text, args.output or 'output.docx')
    else:
        default_folder = r'"E:\czh要的程序\arknights_dialogue\主线剧情一览\靶向药物"'
        if os.path.isdir(default_folder):
            print(f'未指定参数，使用默认文件夹: {default_folder}')
            batch_generate(default_folder)
        else:
            print('未指定输入。用法:')
            print('  批量模式: python generate_docx.py -f "文件夹路径" -o output.docx')
            print('  JSON模式: python generate_docx.py -j data.json -o output.docx')
            print('  文本模式: python generate_docx.py -t source.txt -o output.docx')
